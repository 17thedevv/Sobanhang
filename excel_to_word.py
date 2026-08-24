import pandas as pd
from docx import Document
from docx.shared import Pt
import sys
import os

def excel_to_word(excel_path, word_path):
    print(f"Đang đọc file Excel: {excel_path}")
    try:
        # Đọc tất cả các sheet trong file Excel
        xls = pd.read_excel(excel_path, sheet_name=None)
    except Exception as e:
        print(f"Lỗi khi đọc file Excel: {e}")
        return

    # Tạo một file Word mới
    doc = Document()
    doc.add_heading('Báo cáo phân tích Excel', 0)

    for sheet_name, df in xls.items():
        print(f"Đang xử lý sheet: {sheet_name}")
        
        # Thêm tiêu đề sheet vào Word
        doc.add_heading(f'Sheet: {sheet_name}', level=1)
        
        # Nếu sheet rỗng thì bỏ qua
        if df.empty:
            doc.add_paragraph("(Sheet rỗng)")
            continue

        # Lấy tên cột
        columns = list(df.columns)
        
        # Thêm một bảng vào Word
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        
        # Viết header cho bảng
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = str(col)
            # Chỉnh font header in đậm
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            
        # Thêm dữ liệu từng dòng vào bảng
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                # Ép kiểu dữ liệu về string, nếu NaN thì để trống
                val = row[col]
                if pd.isna(val):
                    val_str = ""
                else:
                    val_str = str(val)
                row_cells[i].text = val_str

        # Thêm dòng trắng ngăn cách giữa các sheet
        doc.add_paragraph("\n")

    # Lưu file Word
    print(f"Đang lưu file Word: {word_path}")
    doc.save(word_path)
    print("Hoàn tất!")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Sử dụng: python excel_to_word.py <đường_dẫn_file_excel> [đường_dẫn_file_word_đầu_ra]")
        print("Ví dụ: python excel_to_word.py docs/QuanLyBanHang_EPIC_1.xlsx docs/Output.docx")
    else:
        input_excel = sys.argv[1]
        
        if len(sys.argv) >= 3:
            output_word = sys.argv[2]
        else:
            # Tự động lấy tên file excel đổi đuôi thành .docx
            base_name = os.path.splitext(input_excel)[0]
            output_word = f"{base_name}_converted.docx"
            
        excel_to_word(input_excel, output_word)
